"""
Text extraction from resume files (PDF and DOCX).

Why this matters: resumes come in wildly inconsistent formats — single column,
two column, tables, text boxes. We extract raw text here and leave the
"making sense of it" job to the LLM extraction step.

CHANGE LOG (v1.1):
- Added logging of extraction attempts, successes, and failures.
- Wrapped library-level errors (corrupted files) into CorruptedFileError
  so the UI can show a friendly message instead of a raw traceback.
- Added EmptyResumeError for files that parse but contain no usable text
  (e.g., scanned image PDFs with no text layer).
- Public function signatures (extract_text, extract_text_from_pdf,
  extract_text_from_docx) are UNCHANGED — existing callers keep working.
"""

import fitz  # PyMuPDF
import docx
from pathlib import Path

from src.config import get_logger, MIN_RESUME_TEXT_LENGTH
from src.exceptions import CorruptedFileError, EmptyResumeError, UnsupportedFileTypeError

logger = get_logger(__name__)


def extract_text_from_pdf(file_path: str) -> str:
    """
    Extract plain text from a PDF file using PyMuPDF.

    PyMuPDF reads text block-by-block, which handles multi-column resumes
    much better than naive string extraction.

    Raises:
        CorruptedFileError: if the PDF can't be opened/parsed at all.
    """
    text_chunks = []
    try:
        with fitz.open(file_path) as doc:
            for page in doc:
                text_chunks.append(page.get_text("text"))
    except Exception as e:
        logger.error(f"Failed to open/parse PDF '{file_path}': {e}")
        raise CorruptedFileError(
            "This PDF couldn't be opened. It may be corrupted, password-protected, "
            "or not a valid PDF file."
        ) from e

    return "\n".join(text_chunks).strip()


def extract_text_from_docx(file_path: str) -> str:
    """
    Extract plain text from a .docx file, including paragraphs and tables.

    Many resumes use tables for layout (e.g., skills in a side column),
    so we explicitly walk both paragraphs and tables.

    Raises:
        CorruptedFileError: if the DOCX can't be opened/parsed at all.
    """
    try:
        document = docx.Document(file_path)
    except Exception as e:
        logger.error(f"Failed to open/parse DOCX '{file_path}': {e}")
        raise CorruptedFileError(
            "This Word document couldn't be opened. It may be corrupted or not a "
            "valid .docx file (note: old .doc files are not supported, only .docx)."
        ) from e

    text_chunks = []

    for para in document.paragraphs:
        if para.text.strip():
            text_chunks.append(para.text)

    for table in document.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                text_chunks.append(row_text)

    return "\n".join(text_chunks).strip()


def extract_text(file_path: str) -> str:
    """
    Dispatch to the right extractor based on file extension, then validate
    that meaningful text was actually extracted.

    Raises:
        UnsupportedFileTypeError: if the extension isn't .pdf or .docx.
        CorruptedFileError: if the file can't be opened/parsed.
        EmptyResumeError: if extraction succeeds but yields little/no text
            (most commonly a scanned image PDF with no real text layer).
    """
    suffix = Path(file_path).suffix.lower()
    logger.info(f"Extracting text from '{Path(file_path).name}' (type={suffix})")

    if suffix == ".pdf":
        text = extract_text_from_pdf(file_path)
    elif suffix == ".docx":
        text = extract_text_from_docx(file_path)
    else:
        logger.warning(f"Rejected unsupported file type: {suffix}")
        raise UnsupportedFileTypeError(f"Unsupported file type: {suffix}. Please upload a PDF or DOCX file.")

    if len(text) < MIN_RESUME_TEXT_LENGTH:
        logger.warning(f"Extracted only {len(text)} characters from '{Path(file_path).name}' — likely empty/scanned.")
        raise EmptyResumeError(
            "We couldn't find readable text in this file. If it's a scanned image "
            "PDF (a photo of a resume rather than a real text document), try "
            "exporting your resume directly from Word/Google Docs as a PDF instead."
        )

    logger.info(f"Successfully extracted {len(text)} characters from '{Path(file_path).name}'")
    return text


if __name__ == "__main__":
    # Quick manual test — run this file directly with a sample resume path
    import sys
    if len(sys.argv) > 1:
        result = extract_text(sys.argv[1])
        print(result)
    else:
        print("Usage: python extractor.py <path_to_resume>")
